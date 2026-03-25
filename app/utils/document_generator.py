"""
Module de génération de documents structurés pour les analyses CPT
Génère des rapports Word, PDF et tableaux formatés pour les réponses expertes
"""

import os
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path

# Document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Table formatting
from tabulate import tabulate
from prettytable import PrettyTable
import markdown
from rich.console import Console
from rich.table import Table as RichTable
from rich.panel import Panel
from rich.text import Text
from rich.columns import Columns

class DocumentGenerator:
    """Classe principale pour la génération de documents structurés"""

    def __init__(self, output_dir: str = "reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.console = Console()

        # Configuration des styles
        self.setup_styles()

    def setup_styles(self):
        """Configure les styles pour les documents"""
        self.word_styles = {
            'title': {'size': 16, 'bold': True, 'color': RGBColor(0, 51, 102)},
            'heading1': {'size': 14, 'bold': True, 'color': RGBColor(0, 102, 204)},
            'heading2': {'size': 12, 'bold': True, 'color': RGBColor(51, 153, 255)},
            'normal': {'size': 11, 'bold': False, 'color': RGBColor(0, 0, 0)},
            'emphasis': {'size': 11, 'bold': True, 'color': RGBColor(204, 0, 0)}
        }

    def format_expert_response_rich(self, response_data: Dict[str, Any]) -> str:
        """Formate une réponse experte avec Rich pour l'affichage console"""
        self.console.print("\n" + "="*80, style="bold blue")
        self.console.print(f"🧪 {response_data.get('title', 'Analyse Expert')}", style="bold cyan")
        self.console.print("="*80, style="bold blue")

        # Sections principales
        for section_name, section_content in response_data.get('sections', {}).items():
            if isinstance(section_content, dict) and 'table' in section_content:
                # Affichage de tableau
                self._display_rich_table(section_name, section_content)
            else:
                # Affichage de texte
                self._display_rich_text(section_name, section_content)

        # Recommandations
        if 'recommendations' in response_data:
            self.console.print("\n💡 RECOMMANDATIONS PRATIQUES", style="bold green")
            for rec in response_data['recommendations']:
                self.console.print(f"  • {rec}", style="green")

        # Conclusion
        if 'conclusion' in response_data:
            self.console.print(f"\n🎯 CONCLUSION", style="bold yellow")
            self.console.print(f"  {response_data['conclusion']}", style="yellow")

        return "Réponse affichée avec formatage Rich"

    def _display_rich_table(self, title: str, table_data: Dict[str, Any]):
        """Affiche un tableau avec Rich"""
        self.console.print(f"\n📊 {title.upper()}", style="bold magenta")

        if 'headers' in table_data and 'rows' in table_data:
            table = RichTable(title=title, show_header=True, header_style="bold magenta")

            # Ajout des colonnes
            for header in table_data['headers']:
                table.add_column(header, style="cyan", justify="center")

            # Ajout des lignes
            for row in table_data['rows']:
                table.add_row(*[str(cell) for cell in row])

            self.console.print(table)

    def _display_rich_text(self, title: str, content: str):
        """Affiche du texte formaté avec Rich"""
        self.console.print(f"\n📋 {title.upper()}", style="bold blue")

        if isinstance(content, list):
            for item in content:
                self.console.print(f"  • {item}", style="white")
        else:
            # Formatage du texte avec des couleurs pour les éléments importants
            formatted_content = content.replace("**", "").replace("*", "")
            self.console.print(f"  {formatted_content}", style="white")

    def generate_word_report(self, response_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Génère un rapport Word structuré"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"analyse_cpt_{timestamp}.docx"

        filepath = self.output_dir / filename
        doc = Document()

        # Titre principal
        title = doc.add_heading(response_data.get('title', 'Rapport d\'Analyse CPT'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")  # Ligne vide

        # Sections
        for section_name, section_content in response_data.get('sections', {}).items():
            heading = doc.add_heading(section_name, level=1)
            self._apply_style_to_paragraph(heading, 'heading1')

            if isinstance(section_content, dict) and 'table' in section_content:
                self._add_table_to_word(doc, section_content)
            else:
                self._add_text_to_word(doc, section_content)

        # Recommandations
        if 'recommendations' in response_data:
            heading = doc.add_heading('Recommandations Pratiques', level=1)
            self._apply_style_to_paragraph(heading, 'heading1')

            for rec in response_data['recommendations']:
                p = doc.add_paragraph(rec, style='List Bullet')
                self._apply_style_to_paragraph(p, 'normal')

        # Conclusion
        if 'conclusion' in response_data:
            heading = doc.add_heading('Conclusion', level=1)
            self._apply_style_to_paragraph(heading, 'heading1')

            p = doc.add_paragraph(response_data['conclusion'])
            self._apply_style_to_paragraph(p, 'emphasis')

        # Sauvegarde
        doc.save(filepath)
        return str(filepath)

    def generate_pdf_report(self, response_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Génère un rapport PDF professionnel"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"analyse_cpt_{timestamp}.pdf"

        filepath = self.output_dir / filename

        # Configuration du document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        ))

        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=20,
            textColor=colors.blue
        ))

        styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12
        ))

        # Contenu
        story = []

        # Titre
        title = Paragraph(response_data.get('title', 'Rapport d\'Analyse CPT'), styles['CustomTitle'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Date
        date_text = f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        story.append(Paragraph(date_text, styles['CustomNormal']))
        story.append(Spacer(1, 20))

        # Sections
        for section_name, section_content in response_data.get('sections', {}).items():
            heading = Paragraph(section_name, styles['CustomHeading1'])
            story.append(heading)

            if isinstance(section_content, dict) and 'table' in section_content:
                table = self._create_pdf_table(section_content)
                story.append(table)
                story.append(Spacer(1, 12))
            else:
                content = self._format_text_for_pdf(section_content)
                story.append(Paragraph(content, styles['CustomNormal']))

            story.append(Spacer(1, 12))

        # Recommandations
        if 'recommendations' in response_data:
            heading = Paragraph('Recommandations Pratiques', styles['CustomHeading1'])
            story.append(heading)

            for rec in response_data['recommendations']:
                story.append(Paragraph(f"• {rec}", styles['CustomNormal']))

            story.append(Spacer(1, 20))

        # Conclusion
        if 'conclusion' in response_data:
            heading = Paragraph('Conclusion', styles['CustomHeading1'])
            story.append(heading)

            story.append(Paragraph(response_data['conclusion'], styles['CustomNormal']))

        # Génération du PDF
        doc.build(story)
        return str(filepath)

    def format_response_with_tables(self, text_response: str) -> str:
        """Améliore une réponse texte avec des tableaux formatés"""
        # Recherche de sections de tableau dans le texte
        lines = text_response.split('\n')
        formatted_lines = []

        in_table = False
        table_lines = []

        for line in lines:
            if '|' in line and ('Type de Sol' in line or 'Valeur' in line or 'Paramètre' in line):
                # Début de tableau
                in_table = True
                table_lines = [line]
            elif in_table and '|' in line:
                # Ligne de tableau
                table_lines.append(line)
            elif in_table and line.strip() == '':
                # Fin de tableau
                in_table = False
                if table_lines:
                    formatted_table = self._format_table_markdown(table_lines)
                    formatted_lines.append(formatted_table)
                    formatted_lines.append('')
                table_lines = []
            elif in_table:
                # Fin de tableau (ligne non vide sans |)
                in_table = False
                if table_lines:
                    formatted_table = self._format_table_markdown(table_lines)
                    formatted_lines.append(formatted_table)
                    formatted_lines.append('')
                table_lines = []
                formatted_lines.append(line)
            else:
                formatted_lines.append(line)

        # Traiter le dernier tableau s'il existe
        if table_lines:
            formatted_table = self._format_table_markdown(table_lines)
            formatted_lines.append(formatted_table)

        return '\n'.join(formatted_lines)

    def _format_table_markdown(self, table_lines: List[str]) -> str:
        """Formate un tableau Markdown en tableau ASCII art"""
        if len(table_lines) < 2:
            return '\n'.join(table_lines)

        # Extraction des données
        headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
        data = []

        for line in table_lines[2:]:  # Skip header and separator
            if '|' in line:
                row = [col.strip() for col in line.split('|')[1:-1]]
                if row:
                    data.append(row)

        # Formatage avec tabulate
        if data:
            return tabulate(data, headers=headers, tablefmt="grid", stralign="center")
        else:
            return '\n'.join(table_lines)

    def _apply_style_to_paragraph(self, paragraph, style_name: str):
        """Applique un style à un paragraphe Word"""
        if style_name in self.word_styles:
            style_config = self.word_styles[style_name]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.size = Pt(style_config['size'])
            run.font.bold = style_config['bold']
            run.font.color.rgb = style_config['color']

    def _add_table_to_word(self, doc: Document, table_data: Dict[str, Any]):
        """Ajoute un tableau à un document Word"""
        if 'headers' in table_data and 'rows' in table_data:
            table = doc.add_table(rows=1, cols=len(table_data['headers']))
            table.style = 'Table Grid'

            # En-têtes
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_data['headers']):
                hdr_cells[i].text = str(header)
                self._apply_style_to_paragraph(hdr_cells[i].paragraphs[0], 'emphasis')

            # Données
            for row_data in table_data['rows']:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)

    def _add_text_to_word(self, doc: Document, content):
        """Ajoute du texte à un document Word"""
        if isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(item, style='List Bullet')
                self._apply_style_to_paragraph(p, 'normal')
        else:
            p = doc.add_paragraph(str(content))
            self._apply_style_to_paragraph(p, 'normal')

    def _create_pdf_table(self, table_data: Dict[str, Any]):
        """Crée un tableau pour PDF"""
        if 'headers' in table_data and 'rows' in table_data:
            data = [table_data['headers']] + table_data['rows']

            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.blue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            return table
        return None

    def _format_text_for_pdf(self, content) -> str:
        """Formate le texte pour PDF"""
        if isinstance(content, list):
            return '<br/>'.join([f"• {item}" for item in content])
        else:
            return str(content).replace('\n', '<br/>')

# Fonctions utilitaires pour l'intégration
def create_structured_response(parameter: str, analysis_data: Dict[str, Any]) -> Dict[str, Any]:
    """Crée une réponse structurée pour un paramètre CPT"""
    response_structure = {
        'title': f'Analyse Expert du paramètre {parameter}',
        'sections': {},
        'recommendations': [],
        'conclusion': ''
    }

    # Structure basée sur le type de paramètre
    if parameter.lower() in ['qc', 'résistance conique']:
        response_structure['sections'] = {
            'Définition Technique': analysis_data.get('definition', ''),
            'Interprétation': {
                'table': {
                    'headers': ['Type de Sol', 'qc (MPa)', 'Classification'],
                    'rows': analysis_data.get('interpretation_table', [])
                }
            },
            'Justifications Techniques': analysis_data.get('justifications', []),
            'Corrélations': analysis_data.get('correlations', '')
        }
        response_structure['recommendations'] = analysis_data.get('recommendations', [])
        response_structure['conclusion'] = analysis_data.get('conclusion', '')

    return response_structure

def generate_reports(response_data: Dict[str, Any], formats: List[str] = ['word', 'pdf']) -> Dict[str, str]:
    """Génère des rapports dans plusieurs formats"""
    generator = DocumentGenerator()
    reports = {}

    if 'word' in formats:
        word_path = generator.generate_word_report(response_data)
        reports['word'] = word_path

    if 'pdf' in formats:
        pdf_path = generator.generate_pdf_report(response_data)
        reports['pdf'] = pdf_path

    return reports